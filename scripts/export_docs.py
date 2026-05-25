"""Export ATLAS planning docs to Office formats.

Generates:
  docs/exports/Nexorus-ATLAS-Architecture.docx    (from the architecture spec markdown)
  docs/exports/Nexorus-ATLAS-Study-Plans.docx     (SOP + register + 22 PM/Arch/QA blocks)
  docs/exports/Nexorus-ATLAS-WBS-Sprint-Plan.xlsx (WBS / sprints / Gantt / risks / milestones)

Both upload cleanly into Google Docs / Google Sheets and open natively in Word / Excel.

Deps (not project deps):  pip install python-docx openpyxl
Usage:                    python scripts/export_docs.py
"""
from __future__ import annotations
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SPEC = os.path.join(ROOT, "docs", "superpowers", "specs",
                    "2026-05-25-atlas-production-architecture-design.md")
OUT = os.path.join(ROOT, "docs", "exports")
os.makedirs(OUT, exist_ok=True)

ACCENT = RGBColor(0x5B, 0x3D, 0xF5)      # nexorus purple
INK = RGBColor(0x1A, 0x1A, 0x2E)


# --------------------------------------------------------------------------- #
#  DOCX: a focused markdown -> docx converter for our spec's constructs
# --------------------------------------------------------------------------- #
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_runs(par, text):
    """Add text to a paragraph, honouring **bold**, `code`, [link](url)."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            r = par.add_run(m.group(1)); r.font.color.rgb = ACCENT
        else:
            par.add_run(tok)


def shade(par, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_fill)
    par._p.get_or_add_pPr().append(shd)


def add_code_block(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.1)
    par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(4)
    shade(par, "F2F2F7")
    for i, line in enumerate(lines):
        # nbsp keeps ASCII-diagram / DDL indentation aligned in Word
        r = par.add_run(line.replace(" ", " "))
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = INK
        if i != len(lines) - 1:
            r.add_break()


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]          # row 1 is the |---| separator
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell.paragraphs[0], "5B3DF5")
    for row in body:
        r = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(r):
                r[j].paragraphs[0].text = ""; add_runs(r[j].paragraphs[0], val)
    doc.add_paragraph()


def new_doc():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    return doc


def render_md(doc, md, first_h1_as_title=True):
    i, first_h1 = 0, first_h1_as_title
    while i < len(md):
        line = md[i]
        if line.startswith("```"):                       # fenced code
            j = i + 1; buf = []
            while j < len(md) and not md[j].startswith("```"):
                buf.append(md[j]); j += 1
            add_code_block(doc, buf); i = j + 1; continue
        if line.startswith("|") and i + 1 < len(md) and set(md[i + 1].strip()) <= set("|-: "):
            j = i
            while j < len(md) and md[j].startswith("|"):
                j += 1
            add_table(doc, md[i:j]); i = j; continue
        if re.match(r"^#{1,4} ", line):                  # heading
            level = len(line) - len(line.lstrip("#")); text = line[level:].strip()
            if level == 1 and first_h1:
                t = doc.add_heading("", level=0); add_runs(t, text); first_h1 = False
            else:
                h = doc.add_heading("", level=min(level, 4)); add_runs(h, text)
        elif line.strip() == "---":
            pass
        elif re.match(r"^\s*[-*] ", line):               # bullet
            indent = len(line) - len(line.lstrip(" "))
            style = ["List Bullet", "List Bullet 2", "List Bullet 3"][min(indent // 2, 2)]
            p = doc.add_paragraph(style=style); add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):              # numbered
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
        elif line.startswith(">"):                       # quote
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(line.lstrip("> ").strip()); r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
        elif line.strip():
            p = doc.add_paragraph(); add_runs(p, line)
        i += 1

    return doc


def build_docx():
    doc = render_md(new_doc(), open(SPEC, encoding="utf-8").read().splitlines())
    path = os.path.join(OUT, "Nexorus-ATLAS-Architecture.docx")
    doc.save(path); return path


STUDY_FILES = [
    "docs/study-plans/README.md",
    "docs/study-plans/atlas/_index.md",
    "docs/study-plans/atlas/0-platform.md",
    "docs/study-plans/atlas/1-watch.md",
    "docs/study-plans/atlas/2-understand.md",
    "docs/study-plans/atlas/3-act.md",
]


def build_study_plans_docx():
    doc = new_doc()
    title = doc.add_heading("", level=0); add_runs(title, "Nexorus ATLAS - Study Plans")
    sub = doc.add_paragraph()
    r = sub.add_run("SOP + register + PM -> Architecture -> QA for all 22 features (v1.0)")
    r.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
    for rel in STUDY_FILES:
        doc.add_page_break()
        render_md(doc, open(os.path.join(ROOT, rel), encoding="utf-8").read().splitlines(),
                  first_h1_as_title=False)
    path = os.path.join(OUT, "Nexorus-ATLAS-Study-Plans.docx")
    doc.save(path); return path


# --------------------------------------------------------------------------- #
#  XLSX: WBS / sprints / Gantt / risks / milestones
# --------------------------------------------------------------------------- #
HDR = Font(bold=True, color="FFFFFF", size=11)
HDR_FILL = PatternFill("solid", fgColor="5B3DF5")
THIN = Border(*[Side(style="thin", color="D9D9E3")] * 4)
WRAP = Alignment(vertical="top", wrap_text=True)
CENTER = Alignment(horizontal="center", vertical="center")
STAGE_FILL = {"0-platform": "EDE7FF", "1-watch": "E2F0FF",
              "2-understand": "E6F7EE", "3-act": "FFF1E0"}
SEV_FILL = {"High": "F8D7DA", "Med": "FFF3CD", "Low": "D4EDDA"}

FEATURES = [
    ("P1", "Monorepo foundation & tooling", "0-platform", "S1", "E1"),
    ("P2", "DigitalOcean infrastructure & CI/CD", "0-platform", "S1", "E1"),
    ("P3", "Database schema, migrations & type generation", "0-platform", "S1-S2", "E2"),
    ("P4", "Object storage (Spaces) integration", "0-platform", "S2", "E2"),
    ("P5", "Authentication - email/password + sessions", "0-platform", "S2", "E3"),
    ("P6", "RBAC, route guards & audit log", "0-platform", "S2", "E3"),
    ("P7", "Observability, hardening, backups & launch", "0-platform", "S1,S6", "E8,E9"),
    ("W1", "Source registry & scheduler", "1-watch", "S3", "E4"),
    ("W2", "RSS & news-API connectors", "1-watch", "S3", "E4"),
    ("W3", "Social connectors (X/IG/FB/TikTok)", "1-watch", "S3-S4", "E4"),
    ("W4", "Normalization, dedup & raw storage", "1-watch", "S3", "E4"),
    ("U1", "LLM provider abstraction & cost ledger", "2-understand", "S4", "E5"),
    ("U2", "Article enrichment (score/issues/sentiment/summary/keywords)", "2-understand", "S4", "E5"),
    ("U3", "Geocoding & incident mapping", "2-understand", "S4", "E5"),
    ("U4", "Crisis snapshots & trends", "2-understand", "S4", "E5"),
    ("U5", "Predictions, insights, actor & leadership analytics", "2-understand", "S4-S5", "E5"),
    ("A1", "Dashboard read API & caching", "3-act", "S2,S5", "E6"),
    ("A2", "Widget integration & live data", "3-act", "S5", "E6"),
    ("A3", "Persisted dashboard layout", "3-act", "S5", "E6"),
    ("A4", "AI assistant - copilot chat", "3-act", "S5", "E7"),
    ("A5", "AI assistant - briefing, forecast & per-widget ask", "3-act", "S5", "E7"),
    ("A6", "Real-time ticker, alerts & War Room", "3-act", "S5-S6", "E8"),
]

SPRINTS = [
    ("S1", "Jun 1-12, 2026", "P1, P2, P3 (start), P7 (skeleton)", "-",
     "Monorepo + DO infra + CI/CD; local stack boots; web deploys to staging."),
    ("S2", "Jun 15-26, 2026", "P3 (finish), P4, P5, P6, A1 (initial)", "M1: DB-backed dashboard",
     "Real login + RBAC; dashboard renders from Postgres, not static JSON."),
    ("S3", "Jun 29-Jul 10, 2026", "W1, W2, W4, W3 (spike)", "-",
     "Scheduled crawl ingests real Indonesian articles; raw to Spaces; dedup."),
    ("S4", "Jul 13-24, 2026", "U1, U2, U3, U4, W3 (cont.)", "M2: live enrichment",
     "Model-agnostic enrichment writes real data; cost tracked; provider switchable."),
    ("S5", "Jul 27-Aug 7, 2026", "U5, A1 (finish), A2, A3, A4, A5, A6 (start)", "M3: feature-complete",
     "All widgets on live data; assistant grounded on DB; layouts persist; realtime."),
    ("S6", "Aug 10-21, 2026", "A6 (finish), P7 (hardening/launch)", "M4: production launch",
     "Security review, load test, backups/DR, runbooks, UAT, prod cutover."),
]

EPICS = [
    ("E1", "Foundation & DevOps", "Monorepo, DO infra (Terraform), CI/CD, Docker Compose, observability skeleton."),
    ("E2", "Data Model & Storage", "Alembic schema (source of truth), seed from static JSON, Kysely codegen, Spaces."),
    ("E3", "Auth & User Management", "Email/password, sessions, RBAC (admin/analyst/viewer), audit log, admin seed."),
    ("E4", "Ingestion Pipeline", "Celery + Beat, SourceConnector, RSS/news/social connectors, dedup, raw store."),
    ("E5", "AI Enrichment Pipeline", "LiteLLM abstraction, scoring/geo/sentiment/predict, snapshots, cost ledger."),
    ("E6", "API/BFF & Frontend Integration", "Read APIs from Postgres, TanStack Query, layouts, trend charts."),
    ("E7", "AI Assistant Productization", "Port assistant to Python ai-api; chat/briefing/forecast/widget; proxy + fallback."),
    ("E8", "Real-time & Observability", "NOTIFY->SSE ticker/alerts, OpenTelemetry, dashboards, freshness alerts."),
    ("E9", "Hardening, Security & Launch", "Security review, load test, backups/DR, runbooks, UAT, cutover."),
]

# Gantt matrix from spec section 15:  2 = primary, 1 = support
GANTT = {
    "E1": {1: 2, 2: 1, 6: 1}, "E2": {1: 1, 2: 2, 3: 1}, "E3": {2: 2, 3: 1},
    "E4": {3: 2, 4: 1}, "E5": {3: 1, 4: 2, 5: 1}, "E6": {2: 1, 4: 1, 5: 2, 6: 1},
    "E7": {4: 1, 5: 2, 6: 1}, "E8": {1: 1, 5: 1, 6: 2}, "E9": {6: 2},
}

RISKS = [
    ("R1", "Social APIs (X/IG/FB/TikTok) can't monitor arbitrary accounts cheaply/legally", "High",
     "Connector abstraction; RSS+news first; paid aggregators per-platform with ToS/legal review + cost cap; social is incremental."),
    ("R2", "LLM enrichment cost scales with article volume", "Med",
     "Cheap model for scoring, strong for summaries; batch; daily budget guardrail; cost ledger."),
    ("R3", "Geocoding accuracy for Indonesian locales", "Med",
     "LLM NER + gazetteer reconcile; unmapped_count surfaced in UI."),
    ("R4", "Scraping / ToS legal exposure", "Med",
     "Prefer official APIs/RSS; legal review before aggregator use; respect robots/ToS."),
    ("R5", "Polyglot monorepo friction (TS + Python)", "Low",
     "Clear boundaries, generated contracts, Taskfile, docker-compose parity."),
    ("R6", "Scope creep into Settings / multi-tenant", "Med",
     "Explicit non-goals; defer to future phase."),
    ("R7", "3-month timeline with 2-3 devs", "Med",
     "Social ingestion incremental; RSS+news+AI is the MVP spine; cut social platforms first if behind."),
]

MILESTONES = [
    ("M1", "DB-backed dashboard (real auth + reads from Postgres)", "End S2", "2026-06-26"),
    ("M2", "Live ingestion + AI enrichment writing real data", "End S4", "2026-07-24"),
    ("M3", "Assistant + real-time on live data (feature-complete)", "End S5", "2026-08-07"),
    ("M4", "Production launch", "End S6", "2026-08-21"),
]


def style_header(ws, ncols, row=1):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = HDR; cell.fill = HDR_FILL; cell.alignment = CENTER; cell.border = THIN


def widths(ws, mapping):
    for col, w in mapping.items():
        ws.column_dimensions[col].width = w


def build_xlsx():
    wb = openpyxl.Workbook()

    # 1) Feature Register ---------------------------------------------------- #
    ws = wb.active; ws.title = "Feature Register"
    cols = ["ID", "Feature", "Stage", "Sprint", "Epic", "Version", "Status"]
    ws.append(cols); style_header(ws, len(cols))
    for fid, name, stage, sprint, epic in FEATURES:
        ws.append([fid, name, stage, sprint, epic, "1.0", "Planned"])
        r = ws.max_row
        ws.cell(r, 3).fill = PatternFill("solid", fgColor=STAGE_FILL[stage])
        for c in range(1, len(cols) + 1):
            ws.cell(r, c).border = THIN; ws.cell(r, c).alignment = WRAP
        ws.cell(r, 1).font = Font(bold=True)
    widths(ws, {"A": 6, "B": 52, "C": 14, "D": 9, "E": 8, "F": 8, "G": 11})
    ws.freeze_panes = "A2"; ws.auto_filter.ref = f"A1:G{ws.max_row}"

    # 2) Sprint Plan --------------------------------------------------------- #
    ws = wb.create_sheet("Sprint Plan")
    cols = ["Sprint", "Window", "Primary features", "Milestone", "Sprint goal / exit"]
    ws.append(cols); style_header(ws, len(cols))
    for s in SPRINTS:
        ws.append(list(s)); r = ws.max_row
        for c in range(1, len(cols) + 1):
            ws.cell(r, c).border = THIN; ws.cell(r, c).alignment = WRAP
        ws.cell(r, 1).font = Font(bold=True)
        if s[3] != "-":
            ws.cell(r, 4).fill = PatternFill("solid", fgColor="E6F7EE")
            ws.cell(r, 4).font = Font(bold=True)
    widths(ws, {"A": 8, "B": 18, "C": 40, "D": 24, "E": 52})
    ws.freeze_panes = "A2"

    # 3) WBS ----------------------------------------------------------------- #
    ws = wb.create_sheet("WBS")
    cols = ["Epic", "Workstream", "Scope"]
    ws.append(cols); style_header(ws, len(cols))
    for e in EPICS:
        ws.append(list(e)); r = ws.max_row
        for c in range(1, len(cols) + 1):
            ws.cell(r, c).border = THIN; ws.cell(r, c).alignment = WRAP
        ws.cell(r, 1).font = Font(bold=True)
    widths(ws, {"A": 7, "B": 34, "C": 80})
    ws.freeze_panes = "A2"

    # 4) Gantt --------------------------------------------------------------- #
    ws = wb.create_sheet("Gantt")
    head = ["Epic", "Workstream"] + [f"{s[0]}\n{s[1].split(',')[0]}" for s in SPRINTS]
    ws.append(head); style_header(ws, len(head))
    primary = PatternFill("solid", fgColor="5B3DF5")
    support = PatternFill("solid", fgColor="C9BEF9")
    for eid, name, _ in EPICS:
        ws.append([eid, name] + [""] * 6); r = ws.max_row
        ws.cell(r, 1).font = Font(bold=True); ws.cell(r, 1).border = THIN
        ws.cell(r, 2).border = THIN; ws.cell(r, 2).alignment = WRAP
        for sidx, val in GANTT[eid].items():
            cell = ws.cell(r, 2 + sidx)
            cell.fill = primary if val == 2 else support
            cell.value = "core" if val == 2 else "support"
            cell.font = Font(color="FFFFFF" if val == 2 else "3A2E7A", size=8, bold=val == 2)
            cell.alignment = CENTER
        for c in range(3, 9):
            ws.cell(r, c).border = THIN
    # milestone row
    ws.append(["", "Milestone"] + ["M1", "", "M2", "", "M3", "M4"])
    # note: place milestones at sprint-end columns
    mrow = ws.max_row
    ws.cell(mrow, 2).font = Font(bold=True, italic=True)
    for c, label in [(4, "M1"), (6, "M2"), (7, "M3"), (8, "M4")]:
        ws.cell(mrow, c).value = label
        ws.cell(mrow, c).fill = PatternFill("solid", fgColor="14B86A")
        ws.cell(mrow, c).font = Font(bold=True, color="FFFFFF"); ws.cell(mrow, c).alignment = CENTER
    for c in [3, 5]:
        ws.cell(mrow, c).value = ""
    widths(ws, {"A": 7, "B": 32})
    for c in range(3, 9):
        ws.column_dimensions[get_column_letter(c)].width = 13
    ws.row_dimensions[1].height = 28; ws.freeze_panes = "C2"

    # legend
    lr = ws.max_row + 2
    ws.cell(lr, 2).value = "Legend:"; ws.cell(lr, 2).font = Font(bold=True)
    ws.cell(lr, 3).value = "core"; ws.cell(lr, 3).fill = primary
    ws.cell(lr, 3).font = Font(color="FFFFFF", bold=True); ws.cell(lr, 3).alignment = CENTER
    ws.cell(lr, 4).value = "support"; ws.cell(lr, 4).fill = support; ws.cell(lr, 4).alignment = CENTER
    ws.cell(lr, 5).value = "milestone"; ws.cell(lr, 5).fill = PatternFill("solid", fgColor="14B86A")
    ws.cell(lr, 5).font = Font(color="FFFFFF", bold=True); ws.cell(lr, 5).alignment = CENTER

    # 5) Risk Register ------------------------------------------------------- #
    ws = wb.create_sheet("Risk Register")
    cols = ["ID", "Risk", "Severity", "Mitigation"]
    ws.append(cols); style_header(ws, len(cols))
    for rid, risk, sev, mit in RISKS:
        ws.append([rid, risk, sev, mit]); r = ws.max_row
        for c in range(1, len(cols) + 1):
            ws.cell(r, c).border = THIN; ws.cell(r, c).alignment = WRAP
        ws.cell(r, 1).font = Font(bold=True)
        ws.cell(r, 3).fill = PatternFill("solid", fgColor=SEV_FILL[sev])
        ws.cell(r, 3).alignment = CENTER
    widths(ws, {"A": 6, "B": 52, "C": 11, "D": 70})
    ws.freeze_panes = "A2"

    # 6) Milestones ---------------------------------------------------------- #
    ws = wb.create_sheet("Milestones")
    cols = ["ID", "Milestone", "When", "Target date"]
    ws.append(cols); style_header(ws, len(cols))
    for m in MILESTONES:
        ws.append(list(m)); r = ws.max_row
        for c in range(1, len(cols) + 1):
            ws.cell(r, c).border = THIN; ws.cell(r, c).alignment = WRAP
        ws.cell(r, 1).font = Font(bold=True, color="14B86A")
    widths(ws, {"A": 6, "B": 54, "C": 10, "D": 14})
    ws.freeze_panes = "A2"

    path = os.path.join(OUT, "Nexorus-ATLAS-WBS-Sprint-Plan.xlsx")
    wb.save(path); return path


def _norm(s):
    for d in ("–", "—", "‒", "―", "−"):
        s = s.replace(d, "-")
    return s.strip()


def verify_consistency():
    """Guard: the workbook's hardcoded feature data must match the markdown index.

    The .docx files are rendered directly from markdown, so they can't drift. The
    .xlsx is built from the FEATURES list above, so we assert it against the single
    source of truth (docs/study-plans/atlas/_index.md). Returns a list of issues.
    """
    idx = os.path.join(ROOT, "docs", "study-plans", "atlas", "_index.md")
    md = {}
    for line in open(idx, encoding="utf-8"):
        m = re.match(r"\|\s*\*\*([PWUA]\d+)\*\*\s*\|(.+)", line)
        if m:
            c = [x.strip() for x in m.group(2).split("|")]
            md[m.group(1)] = (_norm(c[0]), _norm(c[1]), _norm(c[2]), _norm(c[3]))
    script = {f[0]: (_norm(f[1]), _norm(f[2]), _norm(f[3]), _norm(f[4])) for f in FEATURES}
    issues = []
    for fid in sorted(set(md) | set(script)):
        if md.get(fid) != script.get(fid):
            issues.append(f"{fid}: index={md.get(fid)} xlsx={script.get(fid)}")
    if len(md) != 22:
        issues.append(f"expected 22 features in _index.md, found {len(md)}")
    return issues


if __name__ == "__main__":
    issues = verify_consistency()
    if issues:
        print("CONSISTENCY CHECK FAILED (xlsx data vs docs/study-plans/atlas/_index.md):")
        for i in issues:
            print("  -", i)
        raise SystemExit(1)
    print(f"consistency check OK: {len(FEATURES)} features match _index.md")

    outs = [build_docx(), build_study_plans_docx()]
    try:
        outs.append(build_xlsx())
    except PermissionError:
        print("WARNING: xlsx is open/locked - skipped rewrite (close Excel to regenerate).")
    for p in outs:
        print(f"{os.path.relpath(p, ROOT)}  ({os.path.getsize(p):,} bytes)")
